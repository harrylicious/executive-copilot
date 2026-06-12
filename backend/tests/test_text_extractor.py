"""Unit tests for the TextExtractor service.

Validates: Requirements 1.1, 1.4, 1.6
"""

import json
from pathlib import Path
from unittest.mock import patch

import pytest

from app.services.text_extractor import TextExtractor


@pytest.fixture
def extractor():
    """Create a TextExtractor instance."""
    return TextExtractor()


@pytest.fixture
def tmp_txt_file(tmp_path):
    """Create a temporary .txt file with known content."""
    content = "Hello, this is plain text content.\nSecond line here."
    file = tmp_path / "sample.txt"
    file.write_text(content, encoding="utf-8")
    return file, content


@pytest.fixture
def tmp_md_file(tmp_path):
    """Create a temporary .md file with known content."""
    content = "# Heading\n\nThis is **bold** and *italic* markdown.\n\n- Item 1\n- Item 2"
    file = tmp_path / "sample.md"
    file.write_text(content, encoding="utf-8")
    return file, content


@pytest.fixture
def tmp_json_file(tmp_path):
    """Create a temporary .json file with known content."""
    data = {"name": "test", "items": [1, 2, 3], "nested": {"key": "value"}}
    file = tmp_path / "sample.json"
    file.write_text(json.dumps(data), encoding="utf-8")
    return file, data


@pytest.fixture
def tmp_docx_file(tmp_path):
    """Create a temporary .docx file with known content."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("First paragraph of the document.")
    doc.add_paragraph("Second paragraph with more content.")
    file = tmp_path / "sample.docx"
    doc.save(str(file))
    return file, ["First paragraph of the document.", "Second paragraph with more content."]


class TestExtractPlaintext:
    """Tests for .txt file extraction."""

    def test_extracts_txt_content(self, extractor, tmp_txt_file):
        """Should extract full text content from a .txt file."""
        file, expected_content = tmp_txt_file
        result = extractor.extract(file, ".txt")
        assert result == expected_content

    def test_extracts_txt_with_format_without_dot(self, extractor, tmp_txt_file):
        """Should handle format string without leading dot."""
        file, expected_content = tmp_txt_file
        result = extractor.extract(file, "txt")
        assert result == expected_content

    def test_extracts_txt_case_insensitive(self, extractor, tmp_txt_file):
        """Should handle uppercase format string."""
        file, expected_content = tmp_txt_file
        result = extractor.extract(file, ".TXT")
        assert result == expected_content

    def test_extracts_empty_txt_file(self, extractor, tmp_path):
        """Should extract empty string from an empty .txt file."""
        file = tmp_path / "empty.txt"
        file.write_text("", encoding="utf-8")
        result = extractor.extract(file, ".txt")
        assert result == ""


class TestExtractMarkdown:
    """Tests for .md file extraction."""

    def test_extracts_md_content(self, extractor, tmp_md_file):
        """Should extract raw markdown content preserving markup."""
        file, expected_content = tmp_md_file
        result = extractor.extract(file, ".md")
        assert result == expected_content

    def test_preserves_markdown_formatting(self, extractor, tmp_md_file):
        """Should preserve markdown syntax characters."""
        file, _ = tmp_md_file
        result = extractor.extract(file, ".md")
        assert "# Heading" in result
        assert "**bold**" in result
        assert "*italic*" in result

    def test_extracts_md_with_format_without_dot(self, extractor, tmp_md_file):
        """Should handle format string without leading dot."""
        file, expected_content = tmp_md_file
        result = extractor.extract(file, "md")
        assert result == expected_content


class TestExtractJson:
    """Tests for .json file extraction."""

    def test_extracts_json_as_pretty_printed_string(self, extractor, tmp_json_file):
        """Should extract JSON content as pretty-printed string."""
        file, data = tmp_json_file
        result = extractor.extract(file, ".json")
        # Should be valid JSON when parsed back
        parsed = json.loads(result)
        assert parsed == data

    def test_json_output_is_indented(self, extractor, tmp_json_file):
        """Should produce indented JSON output."""
        file, data = tmp_json_file
        result = extractor.extract(file, ".json")
        expected = json.dumps(data, indent=2, ensure_ascii=False)
        assert result == expected

    def test_extracts_json_with_unicode(self, extractor, tmp_path):
        """Should handle JSON with unicode characters."""
        data = {"message": "こんにちは世界", "emoji": "🎉"}
        file = tmp_path / "unicode.json"
        file.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
        result = extractor.extract(file, ".json")
        parsed = json.loads(result)
        assert parsed == data


class TestExtractDocx:
    """Tests for .docx file extraction."""

    def test_extracts_docx_paragraphs(self, extractor, tmp_docx_file):
        """Should extract all paragraphs from a .docx file."""
        file, paragraphs = tmp_docx_file
        result = extractor.extract(file, ".docx")
        for paragraph in paragraphs:
            assert paragraph in result

    def test_docx_paragraphs_joined_with_newlines(self, extractor, tmp_docx_file):
        """Should join paragraphs with newline characters."""
        file, paragraphs = tmp_docx_file
        result = extractor.extract(file, ".docx")
        assert result == "\n".join(paragraphs)

    def test_extracts_docx_with_format_without_dot(self, extractor, tmp_docx_file):
        """Should handle format string without leading dot."""
        file, paragraphs = tmp_docx_file
        result = extractor.extract(file, "docx")
        assert result == "\n".join(paragraphs)


class TestUnsupportedFormats:
    """Tests for unsupported file formats returning None."""

    @pytest.mark.parametrize("fmt", [".png", ".exe", ".html", ".xml", ".zip", ".rtf"])
    def test_unsupported_format_returns_none(self, extractor, tmp_path, fmt):
        """Should return None for unsupported file formats."""
        file = tmp_path / f"file{fmt}"
        file.write_text("some content", encoding="utf-8")
        result = extractor.extract(file, fmt)
        assert result is None

    def test_unsupported_format_without_dot_returns_none(self, extractor, tmp_path):
        """Should return None for unsupported format without leading dot."""
        file = tmp_path / "file.xyz"
        file.write_text("some content", encoding="utf-8")
        result = extractor.extract(file, "xyz")
        assert result is None

    def test_unknown_extension_returns_none(self, extractor, tmp_path):
        """Should return None for completely unknown extensions."""
        file = tmp_path / "file.xyz"
        file.write_text("some content", encoding="utf-8")
        result = extractor.extract(file, ".xyz")
        assert result is None


class TestExtractionFailureHandling:
    """Tests for extraction failure scenarios."""

    def test_nonexistent_file_returns_none(self, extractor, tmp_path):
        """Should return None when file does not exist."""
        file = tmp_path / "nonexistent.txt"
        result = extractor.extract(file, ".txt")
        assert result is None

    def test_invalid_json_returns_none(self, extractor, tmp_path):
        """Should return None when JSON file contains invalid JSON."""
        file = tmp_path / "invalid.json"
        file.write_text("{ not valid json }", encoding="utf-8")
        result = extractor.extract(file, ".json")
        assert result is None

    def test_corrupted_docx_returns_none(self, extractor, tmp_path):
        """Should return None when .docx file is corrupted."""
        file = tmp_path / "corrupted.docx"
        file.write_bytes(b"this is not a valid docx file")
        result = extractor.extract(file, ".docx")
        assert result is None

    def test_permission_error_returns_none(self, extractor, tmp_path):
        """Should return None when file read raises an exception."""
        file = tmp_path / "restricted.txt"
        file.write_text("content", encoding="utf-8")
        with patch.object(Path, "read_text", side_effect=PermissionError("Access denied")):
            result = extractor.extract(file, ".txt")
            assert result is None

    def test_extraction_failure_logs_error(self, extractor, tmp_path, caplog):
        """Should log an error when extraction fails."""
        file = tmp_path / "nonexistent.txt"
        import logging
        with caplog.at_level(logging.ERROR):
            extractor.extract(file, ".txt")
        assert any("Failed to extract text" in record.message for record in caplog.records)

    def test_unsupported_format_logs_warning(self, extractor, tmp_path, caplog):
        """Should log a warning for unsupported formats."""
        file = tmp_path / "file.xyz"
        file.write_text("content", encoding="utf-8")
        import logging
        with caplog.at_level(logging.WARNING):
            extractor.extract(file, ".xyz")
        assert any("Unsupported file format" in record.message for record in caplog.records)
